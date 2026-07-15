"""Document Processor - 文本提取与智能分块

不依赖大型外部库（如 pdfplumber/python-docx 等），实现基础可用的版本：
- .txt / .md: 直接读取
- .pdf: 使用基本的文本提取（若有库可用则尝试，否则读原始字节）
- 其他格式: 尝试作为纯文本处理
"""
from __future__ import annotations

import os
import re
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple


SUPPORTED_EXTENSIONS = {
    ".txt", ".text", ".md", ".markdown",
    ".pdf", ".doc", ".docx", ".rtf",
    ".html", ".htm", ".json", ".csv",
    ".log",
}


class DocumentProcessor:
    """文档处理器 - 负责文本提取和分块"""

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    # ============================================================
    # 文件类型检测
    # ============================================================
    @staticmethod
    def detect_file_type(filepath: str) -> str:
        """根据扩展名检测文件类型"""
        ext = Path(filepath).suffix.lower()
        type_map = {
            ".txt": "text",
            ".text": "text",
            ".md": "markdown",
            ".markdown": "markdown",
            ".pdf": "pdf",
            ".doc": "word",
            ".docx": "word",
            ".rtf": "rich_text",
            ".html": "html",
            ".htm": "html",
            ".json": "json",
            ".csv": "csv",
            ".log": "log",
        }
        return type_map.get(ext, "unknown")

    @staticmethod
    def is_supported(filepath: str) -> bool:
        """检查文件是否支持处理"""
        return Path(filepath).suffix.lower() in SUPPORTED_EXTENSIONS

    # ============================================================
    # 文本提取
    # ============================================================
    def extract_text(self, filepath: str, file_type: Optional[str] = None) -> str:
        """从文件中提取纯文本内容"""
        if not os.path.isfile(filepath):
            raise FileNotFoundError("文件不存在: %s" % filepath)

        if file_type is None:
            file_type = self.detect_file_type(filepath)

        if file_type in ("text", "markdown", "log"):
            return self._extract_text_plain(filepath)
        elif file_type == "pdf":
            return self._extract_text_pdf(filepath)
        elif file_type == "word":
            return self._extract_text_word(filepath)
        elif file_type == "html":
            return self._extract_text_html(filepath)
        elif file_type == "json":
            return self._extract_text_json(filepath)
        elif file_type == "csv":
            return self._extract_text_csv(filepath)
        else:
            return self._extract_text_plain(filepath)

    def _extract_text_plain(self, filepath: str) -> str:
        """处理纯文本文件"""
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            with open(filepath, "r", encoding="gbk", errors="ignore") as f:
                content = f.read()
        return self._clean_text(content)

    def _extract_text_pdf(self, filepath: str) -> str:
        """处理 PDF —— 先尝试 pypdf/pdfplumber，否则回退"""
        # 尝试用 pypdf 读取
        try:
            from pypdf import PdfReader
            reader = PdfReader(filepath)
            text_parts = []
            for page in reader.pages:
                try:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                except Exception:
                    continue
            if text_parts:
                return self._clean_text("\n".join(text_parts))
        except Exception:
            pass

        # 尝试 pdfplumber
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    try:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                    except Exception:
                        continue
            if text_parts:
                return self._clean_text("\n".join(text_parts))
        except Exception:
            pass

        # 回退: 从 PDF 中提取可读文本（去除 PDF 二进制标记）
        try:
            with open(filepath, "rb") as f:
                raw = f.read()
            text = raw.decode("utf-8", errors="ignore")
            # 去除 PDF 操作符和二进制标记
            text = re.sub(r"[A-Za-z]+\s*\[.*?\]\s*TJ?", "", text)
            text = re.sub(r"<<.*?>>", "", text, flags=re.DOTALL)
            text = re.sub(r"/[A-Za-z0-9_]+", "", text)
            return self._clean_text(text)
        except Exception:
            return "[PDF 文件 - 无法自动提取文本，请先转换为纯文本格式]"

    def _extract_text_word(self, filepath: str) -> str:
        """处理 Word 文档 —— 先尝试 python-docx，否则回退"""
        try:
            import docx
            doc = docx.Document(filepath)
            text_parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]

            # 读取表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            if text_parts:
                return self._clean_text("\n".join(text_parts))
        except Exception:
            pass

        return "[Word 文档 - 请先转换为纯文本格式后再上传]"

    def _extract_text_html(self, filepath: str) -> str:
        """处理 HTML —— 先尝试 BeautifulSoup，否则用正则"""
        try:
            from bs4 import BeautifulSoup
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                soup = BeautifulSoup(f.read(), "html.parser")
            for script in soup(["script", "style"]):
                script.decompose()
            text = soup.get_text(separator="\n")
            return self._clean_text(text)
        except Exception:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            text = re.sub(r"<[^>]+>", " ", content)
            return self._clean_text(text)

    def _extract_text_json(self, filepath: str) -> str:
        """处理 JSON —— 格式化后作为文本"""
        import json
        try:
            with open(filepath, "r", encoding="utf-8") as f:
                data = json.load(f)
            return self._clean_text(json.dumps(data, ensure_ascii=False, indent=2))
        except Exception:
            return self._extract_text_plain(filepath)

    def _extract_text_csv(self, filepath: str) -> str:
        """处理 CSV —— 格式化后作为文本"""
        import csv
        try:
            rows = []
            with open(filepath, "r", encoding="utf-8", newline="") as f:
                reader = csv.reader(f)
                for row in reader:
                    if any(cell.strip() for cell in row):
                        rows.append(" | ".join(row))
            return self._clean_text("\n".join(rows))
        except Exception:
            return self._extract_text_plain(filepath)

    def _clean_text(self, text: str) -> str:
        """清理文本 - 去除多余空白、特殊字符等"""
        if not text:
            return ""

        # 统一换行符
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        # 去除控制字符（保留换行、制表符）
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)

        # 合并过多空行（保留段落分隔）
        text = re.sub(r"\n{3,}", "\n\n", text)

        # 去除行尾空格
        text = re.sub(r"[ \t]+\n", "\n", text)

        # 合并空格
        text = re.sub(r" {2,}", " ", text)

        return text.strip()

    # ============================================================
    # 智能分块
    # ============================================================
    def split_chunks(
        self,
        text: str,
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None,
    ) -> List[Dict[str, Any]]:
        """智能分块

        策略:
        1. 先按段落分割
        2. 短段落合并，长段落按句子切分
        3. 保留 chunk_overlap 重叠

        返回: [{"index": 0, "content": "...", "char_count": N}, ...]
        """
        size = chunk_size if chunk_size is not None else self.chunk_size
        overlap = chunk_overlap if chunk_overlap is not None else self.chunk_overlap

        if overlap >= size:
            overlap = size // 4

        text = text.strip()
        if not text:
            return []

        paragraphs = self._split_paragraphs(text)
        segments = []

        for para in paragraphs:
            if len(para) <= size:
                segments.append(para)
            else:
                # 长段落按句子切分
                sentences = self._split_sentences(para)
                segments.extend(sentences)

        # 合并为 chunk
        chunks = []
        current_chunk_parts: List[str] = []
        current_length = 0

        for seg in segments:
            seg_len = len(seg)

            if current_length + seg_len + 1 <= size:
                current_chunk_parts.append(seg)
                current_length += seg_len + 1
            else:
                # 保存当前 chunk
                if current_chunk_parts:
                    chunk_text = " ".join(current_chunk_parts).strip()
                    if chunk_text:
                        chunks.append(chunk_text)

                # 处理重叠
                overlap_parts: List[str] = []
                overlap_length = 0
                for p in reversed(current_chunk_parts):
                    if overlap_length + len(p) + 1 <= overlap:
                        overlap_parts.insert(0, p)
                        overlap_length += len(p) + 1
                    else:
                        break

                current_chunk_parts = list(overlap_parts)
                current_length = overlap_length

                if seg_len <= size:
                    current_chunk_parts.append(seg)
                    current_length += seg_len + 1
                else:
                    # 超长句子直接切块
                    sub_parts = self._split_long_text(seg, size, overlap)
                    for sp in sub_parts:
                        chunks.append(sp)
                    current_chunk_parts = []
                    current_length = 0

        # 最后一个 chunk
        if current_chunk_parts:
            chunk_text = " ".join(current_chunk_parts).strip()
            if chunk_text:
                chunks.append(chunk_text)

        result = []
        for idx, content in enumerate(chunks):
            result.append({
                "index": idx,
                "content": content,
                "char_count": len(content),
            })

        return result

    def _split_paragraphs(self, text: str) -> List[str]:
        """按段落分割"""
        raw_paragraphs = re.split(r"\n\s*\n", text)
        paragraphs = []
        for p in raw_paragraphs:
            p = p.strip()
            if p:
                paragraphs.append(p)
        return paragraphs

    def _split_sentences(self, text: str) -> List[str]:
        """按句子分割（中英文混合）"""
        # 按中英文句号/问号/感叹号/分号切分
        parts = re.split(r"(?<=[。！？.!?;；])\s*", text)
        sentences = []
        for p in parts:
            p = p.strip()
            if p:
                sentences.append(p)
        return sentences

    def _split_long_text(self, text: str, size: int, overlap: int) -> List[str]:
        """暴力切分超长文本（备用）"""
        chunks = []
        start = 0
        n = len(text)
        while start < n:
            end = min(start + size, n)
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            start += size - overlap
        return chunks


__all__ = ["DocumentProcessor", "SUPPORTED_EXTENSIONS"]